from __future__ import annotations

import datetime as dt
import json
from io import BytesIO
from typing import List, Optional

import requests
from fastapi import APIRouter, Depends, HTTPException, Response
from pydantic import BaseModel, EmailStr, Field
from docx import Document

from sqlalchemy.orm import Session

from ..auth import Principal, get_db, get_principal, require_admin
from ..models import AdminUser, User, AuthCredential
from ..security import hash_password


router = APIRouter(prefix="/v1", tags=["users"])


class UserOut(BaseModel):
    id: str
    email: Optional[str] = None
    name: Optional[str] = None
    first_name: Optional[str] = None
    last_name: Optional[str] = None
    dob: Optional[str] = None
    phone: Optional[str] = None
    location: Optional[str] = None
    linkedin_url: Optional[str] = None
    github_url: Optional[str] = None
    portfolio_url: Optional[str] = None

    class Config:
        from_attributes = True


class UserProfilePatch(BaseModel):
    first_name: Optional[str] = None
    last_name: Optional[str] = None
    dob: Optional[str] = None
    phone: Optional[str] = None
    location: Optional[str] = None
    linkedin_url: Optional[str] = None
    github_url: Optional[str] = None
    portfolio_url: Optional[str] = None


@router.get("/users")
def list_users(
    db: Session = Depends(get_db),
    principal: Principal = Depends(get_principal),
):
    if principal.type == "admin":
        rows = (
            db.query(User)
            .join(AdminUser, AdminUser.user_id == User.id)
            .filter(AdminUser.admin_id == principal.id)
            .order_by(User.created_at.desc())
            .all()
        )

        return {"items": [UserOut.model_validate(u).model_dump() for u in rows]}
    else:
        u = db.query(User).filter(User.id == principal.id).first()
        if not u:
            return {"items": []}
        return {"items": [UserOut.model_validate(u).model_dump()]}


@router.patch("/users/{user_id}/profile")
def update_user_profile(
    user_id: str,
    payload: UserProfilePatch,
    db: Session = Depends(get_db),
    principal: Principal = Depends(get_principal),
):
    if principal.type == "user":
        if principal.id != user_id:
            raise HTTPException(status_code=403, detail="Forbidden")
    else:
        if (
            db.query(AdminUser)
            .filter(AdminUser.admin_id == principal.id, AdminUser.user_id == user_id)
            .first()
            is None
        ):
            raise HTTPException(status_code=403, detail="Forbidden")

    user = db.query(User).filter(User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    updates = payload.model_dump(exclude_unset=True)
    for field, value in updates.items():
        setattr(user, field, value)

    if any(k in updates for k in ("first_name", "last_name")):
        user.name = (
            f"{user.first_name or ''} {user.last_name or ''}".strip() or user.name
        )
    user.updated_at = dt.datetime.now()
    db.commit()
    db.refresh(user)
    return {"user": UserOut.model_validate(user).model_dump()}


class AdminCreateUserIn(BaseModel):
    email: EmailStr
    password: str
    firstname: Optional[str] = None
    lastname: Optional[str] = None
    dob: Optional[str] = None
    # optionally link to additional admins too (besides the creator)
    admin_ids: Optional[List[str]] = None


@router.post("/admin/users")
def admin_create_user_and_link(
    payload: AdminCreateUserIn,
    db: Session = Depends(get_db),
    principal: Principal = Depends(get_principal),
):
    require_admin(principal)

    # email unique constraint in AuthCredential
    existing = (
        db.query(AuthCredential)
        .filter(AuthCredential.email == payload.email.lower())
        .first()
    )
    if existing:
        raise HTTPException(status_code=409, detail="Email already registered")

    now = dt.datetime.now()
    user_id = f"u{now.strftime('%Y%m%d%H%M%S')}{now.microsecond}"

    user = User(
        id=user_id,
        name=f"{payload.firstname or ''} {payload.lastname or ''}".strip() or None,
        first_name=payload.firstname,
        last_name=payload.lastname,
        dob=payload.dob,
        created_at=now,
        updated_at=now,
    )
    db.add(user)

    cred = AuthCredential(
        email=payload.email.lower(),
        password_hash=hash_password(payload.password),
        principal_type="user",
        principal_id=user_id,
        principal_name=user.first_name + " " + user.last_name,
        created_at=now,
    )
    db.add(cred)

    # link to creator admin + any others
    admin_ids = set(payload.admin_ids or [])
    admin_ids.add(principal.id)
    for aid in admin_ids:
        db.add(AdminUser(admin_id=aid, user_id=user_id, created_at=now))

    db.commit()
    return {"user": UserOut.model_validate(user).model_dump()}


from ..models import BaseResume
from ..models import StoredFile
from ..storage import save_bytes, safe_filename
from ..resume_docx import extract_resume_json_from_docx
from ..services.pdf_service import docx_bytes_to_pdf_bytes

from fastapi import UploadFile, File


class BaseResumeIn(BaseModel):
    content_text: str = Field(..., min_length=20)


def _ensure_user_access(db: Session, principal: Principal, user_id: str) -> None:
    if principal.type == "user":
        if principal.id != user_id:
            raise HTTPException(status_code=403, detail="Forbidden")
        return
    if (
        db.query(AdminUser)
        .filter(AdminUser.admin_id == principal.id, AdminUser.user_id == user_id)
        .first()
        is None
    ):
        raise HTTPException(status_code=403, detail="Forbidden")


def _get_assigned_resume_template(db: Session, user_id: str) -> StoredFile | None:
    return (
        db.query(StoredFile)
        .filter(
            StoredFile.user_id == user_id,
            StoredFile.application_id == "base",
            StoredFile.kind == "resume_template_docx",
        )
        .order_by(StoredFile.created_at.desc())
        .first()
    )


def _extract_template_preview(sf: StoredFile) -> dict:
    try:
        resp = requests.get(sf.path, timeout=30)
        resp.raise_for_status()
        doc = Document(BytesIO(resp.content))
    except Exception:
        raise HTTPException(
            status_code=500, detail="Failed to read assigned DOCX template"
        )

    lines: list[str] = []
    for paragraph in doc.paragraphs:
        text = " ".join((paragraph.text or "").split()).strip()
        if text:
            lines.append(text)
        if len(lines) >= 12:
            break

    if len(lines) < 12:
        for table in doc.tables:
            for row in table.rows:
                values = []
                for cell in row.cells:
                    text = " ".join(cell.text.split()).strip()
                    if text:
                        values.append(text)
                if values:
                    lines.append(" | ".join(values))
                if len(lines) >= 12:
                    break
            if len(lines) >= 12:
                break

    return {
        "filename": sf.filename,
        "line_count": len(lines),
        "lines": lines,
    }


def _read_stored_docx(sf: StoredFile) -> bytes:
    try:
        resp = requests.get(sf.path, timeout=30)
        resp.raise_for_status()
        return resp.content
    except Exception:
        raise HTTPException(
            status_code=500, detail="Failed to read assigned DOCX template"
        )


@router.put("/users/{user_id}/resume-template-docx")
async def put_resume_template_docx(
    user_id: str,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    principal: Principal = Depends(get_principal),
):
    _ensure_user_access(db, principal, user_id)

    if not (file.filename or "").lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Only .docx files are supported")

    data = await file.read()
    if not data or len(data) < 1000:
        raise HTTPException(status_code=400, detail="Invalid DOCX")

    now = dt.datetime.now()
    filename = safe_filename(file.filename) or "resume_template.docx"
    path = save_bytes(user_id, "base", filename, data)
    sf = StoredFile(
        id=f"resume_template_{user_id}_{now.strftime('%Y%m%d%H%M%S')}{now.microsecond}",
        user_id=user_id,
        application_id="base",
        kind="resume_template_docx",
        path=path,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename,
        created_at=now,
    )
    db.add(sf)
    db.commit()
    return {
        "ok": True,
        "user_id": user_id,
        "stored_file_id": sf.id,
        "filename": sf.filename,
        "updated_at": now.isoformat(),
    }


@router.get("/users/{user_id}/resume-template-preview")
def get_resume_template_preview(
    user_id: str,
    db: Session = Depends(get_db),
    principal: Principal = Depends(get_principal),
):
    _ensure_user_access(db, principal, user_id)
    sf = _get_assigned_resume_template(db, user_id)
    if not sf:
        raise HTTPException(status_code=404, detail="Resume template not found")
    return {
        "user_id": user_id,
        "stored_file_id": sf.id,
        **_extract_template_preview(sf),
    }


@router.get("/users/{user_id}/resume-template-preview.pdf")
def get_resume_template_preview_pdf(
    user_id: str,
    db: Session = Depends(get_db),
    principal: Principal = Depends(get_principal),
):
    _ensure_user_access(db, principal, user_id)
    sf = _get_assigned_resume_template(db, user_id)
    if not sf:
        raise HTTPException(status_code=404, detail="Resume template not found")

    docx_bytes = _read_stored_docx(sf)
    try:
        pdf_bytes = docx_bytes_to_pdf_bytes(docx_bytes)
    except Exception:
        pass
        # raise HTTPException(
        #     status_code=500,
        #     detail="Failed to generate PDF preview for the assigned template",
        # )

    return Response(content=pdf_bytes, media_type="application/pdf")


@router.put("/users/{user_id}/base-resume-docx")
async def put_base_resume_docx(
    user_id: str,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    principal: Principal = Depends(get_principal),
):
    """Upload a base resume DOCX.

    - Extract bullet blocks and store JSON into base_resumes.content_text.
    - Save the original DOCX as a StoredFile(kind='base_resume_docx') so we can export later
      while keeping the original template/format.
    """

    _ensure_user_access(db, principal, user_id)

    if not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Only .docx files are supported")

    data = await file.read()
    if not data or len(data) < 1000:
        raise HTTPException(status_code=400, detail="Invalid DOCX")

    # Extract JSON representation
    try:
        resume_json = extract_resume_json_from_docx(data)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to parse DOCX: {str(e)}")

    now = dt.datetime.now()

    # Upsert BaseResume row
    row = db.get(BaseResume, user_id)
    content_text = json.dumps(resume_json, ensure_ascii=False)
    if row:
        row.content_text = content_text
        row.updated_at = now
    else:
        row = BaseResume(
            user_id=user_id,
            content_text=content_text,
            created_at=now,
            updated_at=now,
        )
        db.add(row)

    # Save original DOCX
    filename = safe_filename(file.filename) or "base_resume.docx"
    path = save_bytes(user_id, "base", filename, data)
    sf = StoredFile(
        id=f"base_resume_{user_id}_{int(now.timestamp())}",
        user_id=user_id,
        application_id="base",
        kind="base_resume_docx",
        path=path,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename,
        created_at=now,
    )
    db.add(sf)

    db.commit()
    exp_count = len(resume_json.get("experiences") or [])
    bullets_count = sum(
        len((e.get("bullets") or [])) for e in (resume_json.get("experiences") or [])
    )
    return {
        "ok": True,
        "user_id": user_id,
        "stored_file_id": sf.id,
        "extracted": {"experiences": exp_count, "bullets": bullets_count},
        "updated_at": now.isoformat(),
    }


@router.put("/users/{user_id}/base-resume")
def put_base_resume(
    user_id: str,
    payload: BaseResumeIn,
    db: Session = Depends(get_db),
    principal: Principal = Depends(get_principal),
):
    _ensure_user_access(db, principal, user_id)

    now = dt.datetime.now()
    row = db.get(BaseResume, user_id)
    if row:
        row.content_text = payload.content_text
        row.updated_at = now
    else:
        row = BaseResume(
            user_id=user_id,
            content_text=payload.content_text,
            created_at=now,
            updated_at=now,
        )
        db.add(row)

    db.commit()
    return {"ok": True, "user_id": user_id, "updated_at": now.isoformat()}


@router.get("/users/{user_id}/base-resume")
def get_base_resume(
    user_id: str,
    db: Session = Depends(get_db),
    principal: Principal = Depends(get_principal),
):
    _ensure_user_access(db, principal, user_id)

    br = db.get(BaseResume, user_id)
    template = _get_assigned_resume_template(db, user_id)
    return {
        "user_id": user_id,
        "content_text": br.content_text if br else "",
        "resume_template_file_id": template.id if template else None,
        "resume_template_filename": template.filename if template else None,
        "resume_template_uploaded": bool(template),
        "updated_at": br.updated_at.isoformat() if br else None,
    }


class BaseResumeListItem(BaseModel):
    user_id: str
    title: str | None = None
    content_text: str = ""
    resume_template_file_id: str | None = None
    resume_template_filename: str | None = None
    resume_template_uploaded: bool = False
    updated_at: str | None = None


@router.get("/base-resumes")
def list_base_resumes(
    db: Session = Depends(get_db),
    principal: Principal = Depends(get_principal),
):
    """List base resumes visible to the current principal.

    - Admin: all linked users (+ their base resume content, if any)
    - User: only themselves
    """
    if principal.type == "admin":
        rows = (
            db.query(User, BaseResume)
            .join(AdminUser, AdminUser.user_id == User.id)
            .outerjoin(BaseResume, BaseResume.user_id == User.id)
            .filter(AdminUser.admin_id == principal.id)
            .order_by(User.created_at.desc())
            .all()
        )
        items = []
        for u, br in rows:
            template = _get_assigned_resume_template(db, u.id)
            items.append(
                {
                    "user_id": u.id,
                    "title": u.name or u.id,
                    "content_text": br.content_text if br else "",
                    "resume_template_file_id": template.id if template else None,
                    "resume_template_filename": template.filename if template else None,
                    "resume_template_uploaded": bool(template),
                    "updated_at": br.updated_at.isoformat() if br else None,
                }
            )
        return {"items": items}
    else:
        u = db.query(User).filter(User.id == principal.id).first()
        br = db.get(BaseResume, principal.id)
        template = _get_assigned_resume_template(db, principal.id)
        return {
            "items": [
                {
                    "user_id": principal.id,
                    "title": (u.name if u else None) or principal.id,
                    "content_text": br.content_text if br else "",
                    "resume_template_file_id": template.id if template else None,
                    "resume_template_filename": template.filename if template else None,
                    "resume_template_uploaded": bool(template),
                    "updated_at": br.updated_at.isoformat() if br else None,
                }
            ]
        }


class LinkUserIn(BaseModel):
    user_id: str


@router.get("/admin/users/all")
def admin_list_all_users(
    db: Session = Depends(get_db),
    principal: Principal = Depends(get_principal),
):
    require_admin(principal)
    rows = db.query(User).order_by(User.created_at.desc()).all()
    return {"items": [UserOut.model_validate(u).model_dump() for u in rows]}


@router.post("/admin/users/link")
def admin_link_user(
    payload: LinkUserIn,
    db: Session = Depends(get_db),
    principal: Principal = Depends(get_principal),
):
    require_admin(principal)

    # verify user exists
    u = db.query(User).filter(User.id == payload.user_id).first()
    if not u:
        raise HTTPException(status_code=404, detail="User not found")

    now = dt.datetime.now()
    existing = (
        db.query(AdminUser)
        .filter(
            AdminUser.admin_id == principal.id, AdminUser.user_id == payload.user_id
        )
        .first()
    )
    if existing:
        return {"ok": True}

    db.add(AdminUser(admin_id=principal.id, user_id=payload.user_id, created_at=now))
    db.commit()
    return {"ok": True}


@router.post("/admin/users/unlink")
def admin_unlink_user(
    payload: LinkUserIn,
    db: Session = Depends(get_db),
    principal: Principal = Depends(get_principal),
):
    require_admin(principal)

    row = (
        db.query(AdminUser)
        .filter(
            AdminUser.admin_id == principal.id, AdminUser.user_id == payload.user_id
        )
        .first()
    )
    if row is None:
        return {"ok": True}

    db.delete(row)
    db.commit()
    return {"ok": True}
