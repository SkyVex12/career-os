from __future__ import annotations

import re
from io import BytesIO
from typing import Any, Dict, List, Optional, Tuple, Set

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


# ---------------------------
# Helpers
# ---------------------------


def _norm_keep_case(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").strip())


def _norm_lower(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").strip().lower())


def _is_bullet_prefix_text(txt: str) -> bool:
    t = (txt or "").lstrip()
    return t.startswith(("•", "-", "*", "\u2022", "‣", "▪", "–", "—", "●"))


def _has_numbering(p: Paragraph) -> bool:
    try:
        pPr = p._p.pPr
        return bool(pPr is not None and pPr.numPr is not None)
    except Exception:
        return False


def _style_looks_like_list(p: Paragraph) -> bool:
    try:
        name = (p.style.name or "").lower()
        return ("list" in name) or ("bullet" in name)
    except Exception:
        return False


def _is_bullet_paragraph(p: Paragraph) -> bool:
    txt = (p.text or "").strip()
    if not txt:
        return False
    # bullets are often true Word lists (numPr) without visible "•"
    return _has_numbering(p) or _style_looks_like_list(p) or _is_bullet_prefix_text(txt)


def _strip_visible_bullet_prefix(t: str) -> str:
    s = (t or "").strip()
    s = re.sub(r"^(?:[•\u2022‣▪●\-\*\–\—]\s+)", "", s)
    return s.strip()


def _looks_like_contact_line(t: str) -> bool:
    """
    Heuristic: skip header/contact lines when finding summary.
    """
    tl = _norm_lower(t)
    if not tl:
        return True
    if "@" in tl:  # email
        return True
    if "linkedin" in tl or "github" in tl or "portfolio" in tl:
        return True
    if re.search(r"\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b", tl):  # phone-ish
        return True
    if re.search(r"\b(city|state|usa|united states|remote)\b", tl):
        return True
    return False


# ---------------------------
# DOCX -> JSON extraction
# ---------------------------


def extract_resume_json_from_docx(docx_bytes: bytes) -> Dict[str, Any]:
    """
    Best-effort DOCX -> structured resume JSON.

    - Extracts summary paragraphs (non-bullets near the top).
    - Extracts bullet blocks as experiences, storing bullet_para_idxs.
    """
    doc = Document(BytesIO(docx_bytes))
    paragraphs = list(doc.paragraphs)

    # ---- 1) Find summary near top ----
    # Strategy:
    # - scan top ~40 paragraphs
    # - skip blank/contact-ish lines
    # - take first "non-bullet content" lines until first bullet block appears
    summary_para_idxs: List[int] = []
    summary_parts: List[str] = []

    seen_any_real_text = False
    for idx, p in enumerate(paragraphs[:40]):
        t = _norm_keep_case(p.text)
        if not t:
            continue

        # mark we are past the name/contact area once we see non-contact line
        if not _looks_like_contact_line(t):
            seen_any_real_text = True

        # stop once bullet section begins (most resumes)
        if _is_bullet_paragraph(p):
            # stop summary capture at first bullet
            if summary_para_idxs:
                break
            continue

        # only start collecting after we passed initial contact-ish lines
        if not seen_any_real_text:
            continue

        # heuristic: summary is usually 1–3 paragraphs, stop if too many
        if len(summary_para_idxs) >= 3:
            break

        # do not treat obvious section headers as summary (optional)
        # if text is very short and all caps, likely a header
        if len(t) <= 18 and t.isupper():
            continue

        summary_para_idxs.append(idx)
        summary_parts.append(t)

    summary_text = " ".join([_norm_keep_case(x) for x in summary_parts]).strip()

    # ---- 2) Extract experience bullet blocks ----
    experiences: List[Dict[str, Any]] = []
    current: Optional[Dict[str, Any]] = None
    prev_nonbullet_text = ""

    for idx, p in enumerate(paragraphs):
        t = _norm_keep_case(p.text)
        if not t:
            if current and current.get("bullets"):
                experiences.append(current)
                current = None
            continue

        if _is_bullet_paragraph(p):
            if current is None:
                current = {
                    "header": prev_nonbullet_text,
                    "company": None,
                    "title": None,
                    "start": None,
                    "end": None,
                    "location": None,
                    "bullets": [],
                    "bullet_para_idxs": [],
                }

            current["bullets"].append(_strip_visible_bullet_prefix(t))
            current["bullet_para_idxs"].append(idx)
        else:
            if current and current.get("bullets"):
                experiences.append(current)
                current = None
            prev_nonbullet_text = t

    if current and current.get("bullets"):
        experiences.append(current)

    return {
        "schema": "resume_json_v2",
        "summary": summary_text,
        "summary_para_idxs": summary_para_idxs,
        "experiences": experiences,
    }


# ---------------------------
# Paragraph ops
# ---------------------------


def _set_paragraph_text_preserve_format(p: Paragraph, text: str) -> None:
    """
    Replace text while preserving paragraph formatting (list/numPr/style).
    Keep first run formatting, remove extra runs.
    """
    text = text or ""
    runs = list(p.runs)
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            try:
                r._r.getparent().remove(r._r)
            except Exception:
                pass
    else:
        p.add_run(text)


# ---------------------------
# Summary replacement
# ---------------------------

ELLIPSIS_RE = re.compile(r"(?:\.{3}|\u2026)\s*$")


def _clean_summary_text(s: str) -> str:
    s = (s or "").strip()
    # If anything ends with ellipsis, remove it (viewer-safe)
    s = ELLIPSIS_RE.sub("", s).strip()
    return s


def _split_to_fit_paragraphs(text: str, n: int) -> List[str]:
    """
    Split text into n parts without creating an overstuffed last paragraph.
    Splits by sentence boundaries first; falls back to word wrapping.
    """
    text = _clean_summary_text(text)
    if n <= 1:
        return [text]

    # Prefer sentence splitting
    sentences = [x.strip() for x in re.split(r"(?<=[.!?])\s+", text) if x.strip()]

    # If not enough sentence structure, fall back to word chunks
    if len(sentences) < 2:
        words = text.split()
        if not words:
            return [""] * n
        # approximate equal chunks
        chunk = max(1, (len(words) + n - 1) // n)
        parts = [" ".join(words[i : i + chunk]) for i in range(0, len(words), chunk)]
        # normalize to exactly n
        parts = (parts + [""] * n)[:n]
        if len(parts) > n:
            parts = parts[: n - 1] + [" ".join(parts[n - 1 :])]
        return parts

    # Greedy distribute sentences into n buckets (keeps lines reasonable)
    buckets = [""] * n
    lengths = [0] * n
    for sent in sentences:
        k = lengths.index(min(lengths))
        buckets[k] = (buckets[k] + " " + sent).strip() if buckets[k] else sent
        lengths[k] += len(sent)

    return buckets


def replace_summary_in_docx(
    docx_bytes: bytes,
    summary_para_idxs: List[int],
    new_summary: str,
) -> bytes:
    """
    Replace summary paragraphs in-place, preserving formatting,
    while avoiding packing everything into one final paragraph (which can render as …).
    """
    print(new_summary)
    doc = Document(BytesIO(docx_bytes))
    paragraphs = list(doc.paragraphs)

    idxs = [
        int(i)
        for i in (summary_para_idxs or [])
        if isinstance(i, int) or str(i).isdigit()
    ]
    idxs = [i for i in idxs if 0 <= i < len(paragraphs)]
    if not idxs:
        return docx_bytes

    # IMPORTANT: don't split on "\n" only; normalize all whitespace, then fit across idxs
    summary_text = re.sub(r"\s+", " ", (new_summary or "")).strip()
    print("++++++", summary_text)
    if not summary_text:
        return docx_bytes

    parts = _split_to_fit_paragraphs(new_summary, n=len(idxs))

    # Write exactly len(idxs) paragraphs
    for j, idx in enumerate(idxs):
        _set_paragraph_text_preserve_format(paragraphs[idx], parts[j])

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


# ---------------------------
# Bullet replacement (your version kept, cleaned)
# ---------------------------


def _tokenize(s: str) -> List[str]:
    s = _norm_lower(s)
    s = re.sub(r"[^a-z0-9+#/.-]+", " ", s)
    return [t for t in s.split(" ") if t]


def _find_target_idxs_by_matching_text(
    paragraphs: List[Paragraph],
    original_bullets: List[str],
) -> List[int]:
    norm_to_idxs: Dict[str, List[int]] = {}
    for i, p in enumerate(paragraphs):
        t = _norm_lower(p.text)
        if not t:
            continue
        norm_to_idxs.setdefault(t, []).append(i)

    target: List[int] = []
    used: Set[int] = set()

    # Exact match
    for b in original_bullets:
        nb = _norm_lower(b)
        found = None
        for idx in norm_to_idxs.get(nb, []):
            if idx not in used:
                found = idx
                break
        if found is not None:
            target.append(found)
            used.add(found)

    # Fuzzy fallback
    if len(target) < len(original_bullets):
        candidates = [
            (i, set(_tokenize(paragraphs[i].text)))
            for i in range(len(paragraphs))
            if _is_bullet_paragraph(paragraphs[i])
        ]
        for b in original_bullets[len(target) :]:
            btoks = set(_tokenize(b))
            best = None
            best_score = 0.0
            for i, ptoks in candidates:
                if i in used:
                    continue
                if not btoks or not ptoks:
                    continue
                overlap = len(btoks & ptoks) / max(1, len(btoks))
                if overlap > best_score:
                    best_score = overlap
                    best = i
            if best is not None and best_score >= 0.70:
                target.append(best)
                used.add(best)

    return sorted(target)


def replace_bullets_in_docx(
    docx_bytes: bytes,
    bullet_blocks: List[Dict[str, Any]],
    new_bullets_by_block_index: Dict[int, List[str]],
) -> bytes:
    doc = Document(BytesIO(docx_bytes))

    def get_paragraphs() -> List[Paragraph]:
        return list(doc.paragraphs)

    work: List[Tuple[int, List[int]]] = []
    for block_idx, block in enumerate(bullet_blocks):
        if block_idx not in new_bullets_by_block_index:
            continue

        raw_idxs = block.get("bullet_para_idxs") or []
        target_idxs: List[int] = []
        for i in raw_idxs:
            try:
                target_idxs.append(int(i))
            except Exception:
                pass

        # Fallback: match by original bullet text if idxs missing
        if not target_idxs:
            original_bullets = [
                str(x).strip() for x in (block.get("bullets") or []) if str(x).strip()
            ]
            if original_bullets:
                target_idxs = _find_target_idxs_by_matching_text(
                    get_paragraphs(), original_bullets
                )

        target_idxs = [i for i in target_idxs if i >= 0]
        if target_idxs:
            work.append((block_idx, sorted(target_idxs)))

    work.sort(key=lambda x: x[1][-1], reverse=True)

    for block_idx, target in work:
        paragraphs = get_paragraphs()
        target = [i for i in target if 0 <= i < len(paragraphs)]
        if not target:
            continue

        new_bullets = [
            (b or "").strip()
            for b in (new_bullets_by_block_index.get(block_idx) or [])
            if (b or "").strip()
        ]
        if not new_bullets:
            continue

        original_texts = [(paragraphs[i].text or "").strip() for i in target]

        # Enforce SAME count as template
        if len(new_bullets) < len(original_texts):
            new_bullets = new_bullets + original_texts[len(new_bullets) :]
        elif len(new_bullets) > len(original_texts):
            new_bullets = new_bullets[: len(original_texts)]

        # Overwrite 1:1
        for j, idx in enumerate(target):
            paragraphs = get_paragraphs()
            if 0 <= idx < len(paragraphs):
                _set_paragraph_text_preserve_format(paragraphs[idx], new_bullets[j])

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


_BOLD_TAG_RE = re.compile(r"(<b>.*?</b>)", re.IGNORECASE | re.DOTALL)
_STRIP_TAG_RE = re.compile(r"</?b>", re.IGNORECASE)


def _normalize_bold_markup(text: str) -> str:
    text = (text or "").replace("<br>", " ").replace("<br/>", " ").strip()
    text = re.sub(r"<\s*/\s*b\s*>", "</b>", text, flags=re.IGNORECASE)
    text = re.sub(r"<\s*b\s*>", "<b>", text, flags=re.IGNORECASE)
    open_count = len(re.findall(r"<b>", text, flags=re.IGNORECASE))
    close_count = len(re.findall(r"</b>", text, flags=re.IGNORECASE))
    if close_count > open_count:
        diff = close_count - open_count
        for _ in range(diff):
            text = re.sub(r"</b>", "", text, count=1, flags=re.IGNORECASE)
    elif open_count > close_count:
        text += "</b>" * (open_count - close_count)
    return text


def _clean_markup_text(text: str) -> str:
    return _normalize_bold_markup(text)


def _add_markup_paragraph(
    doc: Document,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    size: int = 10,
    align: WD_ALIGN_PARAGRAPH | None = None,
    left_indent: float = 0.0,
    space_after: int = 0,
) -> Paragraph:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.left_indent = Inches(left_indent)
    p.paragraph_format.space_after = Pt(space_after)

    parts = _BOLD_TAG_RE.split(_clean_markup_text(text))
    for part in parts:
        if not part:
            continue
        is_bold_part = part.lower().startswith("<b>") and part.lower().endswith("</b>")
        run = p.add_run(_STRIP_TAG_RE.sub("", part))
        run.bold = bold or is_bold_part
        run.italic = italic
        run.font.size = Pt(size)
        run.font.name = "Calibri"
    return p


def _add_hyperlink(paragraph: Paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_contact_line(doc: Document, items: list) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    first = True
    for item in items:
        if isinstance(item, dict):
            text = str(item.get("text") or item.get("label") or "").strip()
            url = str(item.get("url") or "").strip()
        else:
            text = str(item).strip()
            url = ""
        if not text:
            continue
        if not first:
            sep = p.add_run(" | ")
            sep.font.size = Pt(9)
            sep.font.name = "Calibri"
        if url:
            _add_hyperlink(p, text, url)
        else:
            run = p.add_run(text)
            run.font.size = Pt(9)
            run.font.name = "Calibri"
        first = False


def _add_experience_header(
    doc: Document,
    *,
    company: str,
    location: str,
    duration: str,
    job_title: str,
) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    left = table.cell(0, 0).paragraphs[0]
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    left_parts = [job_title, company]
    if location:
        left_parts.append(location)
    company_line = " | ".join([part for part in left_parts if part])
    company_run = left.add_run(company_line)
    company_run.bold = True
    company_run.font.name = "Calibri"
    company_run.font.size = Pt(10.5)

    right_run = right.add_run(duration or "")
    right_run.font.name = "Calibri"
    right_run.font.size = Pt(10)


def build_resume_docx_bytes(resume: Dict[str, Any]) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    job_title = (resume.get("job_title") or "Software Engineer").strip()
    candidate = resume.get("candidate") or {}
    candidate_name = (candidate.get("name") or "").strip()
    contact_items = candidate.get("contact_items") or []
    summary = (resume.get("summary") or "").strip()
    skills = resume.get("skills") or []
    experiences = resume.get("experiences") or []
    education = resume.get("education") or []

    if candidate_name:
        _add_markup_paragraph(
            doc,
            candidate_name.upper(),
            bold=True,
            size=18,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=4,
        )
    if contact_items:
        _add_contact_line(doc, contact_items)
    _add_markup_paragraph(
        doc,
        job_title,
        bold=True,
        size=13,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=8,
    )

    sections = [
        ("Summary", summary),
        ("Skills", skills),
        ("Experience", experiences),
        ("Education", education),
    ]

    for idx, (title, payload) in enumerate(sections):
        _add_markup_paragraph(doc, title.upper(), bold=True, size=11.5, space_after=4)

        if title == "Summary":
            _add_markup_paragraph(doc, str(payload or ""), size=10, space_after=8)

        elif title == "Skills":
            for item in payload:
                category = str(item.get("category") or "").strip()
                values = [str(v).strip() for v in (item.get("items") or []) if str(v).strip()]
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.1)
                p.paragraph_format.space_after = Pt(2)
                prefix = p.add_run("- ")
                prefix.font.size = Pt(10)
                prefix.font.name = "Calibri"
                if category:
                    cat = p.add_run(category)
                    cat.bold = True
                    cat.font.size = Pt(10)
                    cat.font.name = "Calibri"
                    colon = p.add_run(": ")
                    colon.font.size = Pt(10)
                    colon.font.name = "Calibri"
                vals = p.add_run(", ".join(values))
                vals.font.size = Pt(10)
                vals.font.name = "Calibri"
            doc.add_paragraph().paragraph_format.space_after = Pt(4)

        elif title == "Experience":
            for exp in payload:
                company = str(exp.get("company") or "").strip()
                location = str(exp.get("location") or "").strip()
                _add_experience_header(
                    doc,
                    company=company,
                    location=location,
                    duration=str(exp.get("duration") or "").strip(),
                    job_title=str(exp.get("job_title") or "").strip(),
                )
                for sentence in exp.get("sentences") or []:
                    _add_markup_paragraph(
                        doc,
                        f"- {_clean_markup_text(str(sentence))}",
                        size=10,
                        left_indent=0.18,
                        space_after=2,
                    )
                doc.add_paragraph().paragraph_format.space_after = Pt(3)

        elif title == "Education":
            for item in payload:
                _add_markup_paragraph(
                    doc,
                    str(item.get("school") or "").strip(),
                    size=10,
                    space_after=1,
                )
                _add_markup_paragraph(
                    doc,
                    str(item.get("degree") or "").strip(),
                    size=10,
                    space_after=1,
                )
                _add_markup_paragraph(
                    doc,
                    str(item.get("duration") or "").strip(),
                    size=10,
                    space_after=3,
                )

        if idx < len(sections) - 1:
            doc.add_paragraph().paragraph_format.space_after = Pt(6)

    out = BytesIO()
    doc.save(out)
    return out.getvalue()
