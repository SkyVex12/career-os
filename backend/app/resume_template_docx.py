from __future__ import annotations

from copy import deepcopy
from io import BytesIO
import re
from typing import Any, Dict, Iterable, List

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run


def render_resume_template_docx_bytes(
    template_bytes: bytes,
    resume: Dict[str, Any],
) -> bytes:
    doc = Document(BytesIO(template_bytes))

    candidate = resume.get("candidate") or {}
    contact_items = candidate.get("contact_items") or []
    experiences = _normalize_experiences(resume.get("experiences") or [])
    skills = _normalize_skills(resume.get("skills") or [])
    education = _normalize_education(resume.get("education") or [])

    replacements = {
        "[Phone_number]": _contact_value(contact_items, "phone"),
        "[phone_number]": _contact_value(contact_items, "phone"),
        "[Email]": _contact_value(contact_items, "email"),
        "[email]": _contact_value(contact_items, "email"),
        "[Adress]": _contact_value(contact_items, "address"),
        "[adress]": _contact_value(contact_items, "address"),
        "[Address]": _contact_value(contact_items, "address"),
        "[Linkedin]": _contact_value(contact_items, "linkedin"),
        "[LinkedIn]": _contact_value(contact_items, "linkedin"),
        "[LINKEDIN]": _contact_value(contact_items, "linkedin"),
        "[Total_role]": str(resume.get("job_title") or "").strip(),
        "[Total Role]": str(resume.get("job_title") or "").strip(),
        "[Skills]": "SKILLS",
        "[SKILLS]": "SKILLS",
        "[Experience]": "EXPERIENCE",
        "[EXPERIENCE]": "EXPERIENCE",
        "[Education]": "EDUCATION",
        "[EDUCATION]": "EDUCATION",
    }

    _render_name_block(doc, str(candidate.get("name") or ""))
    _render_contact_link_block(doc, ["[Linkedin]", "[LinkedIn]", "[LINKEDIN]"], _contact_item(contact_items, "linkedin"))
    _render_heading_block(doc, "[Education]", "EDUCATION")
    _render_heading_block(doc, "[EDUCATION]", "EDUCATION")
    _replace_tokens_in_document(doc, replacements)
    _render_summary_block(doc, str(resume.get("summary") or ""))
    _render_skill_block(doc, skills)
    _render_experience_block(doc, experiences)
    _render_education_block(doc, education)
    _clear_leftover_placeholders(doc)

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def _normalize_experiences(items: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    out: List[Dict[str, Any]] = []
    for item in items:
        bullets = item.get("sentences") or item.get("bullets") or []
        out.append(
            {
                "role": str(item.get("job_title") or item.get("title") or "").strip(),
                "company": str(item.get("company") or "").strip(),
                "location": str(item.get("location") or "").strip(),
                "duration": str(item.get("duration") or "").strip(),
                "bullets": [_normalize_bullet_text(str(x)) for x in bullets if str(x).strip()],
            }
        )
    return out


def _normalize_skills(items: List[Dict[str, Any]]) -> List[Dict[str, str]]:
    out: List[Dict[str, str]] = []
    for item in items:
        category = str(item.get("category") or "").strip()
        values = [str(v).strip() for v in (item.get("items") or []) if str(v).strip()]
        details = ", ".join(values)
        if category or details:
            out.append({"category": category, "details": details})
    return out


def _normalize_education(items: List[Dict[str, Any]]) -> List[Dict[str, str]]:
    out: List[Dict[str, str]] = []
    for item in items:
        out.append(
            {
                "school": str(item.get("school") or "").strip(),
                "degree": str(item.get("degree") or "").strip(),
                "duration": str(item.get("duration") or "").strip(),
            }
        )
    return out


def _contact_text(items: List[Any], index: int) -> str:
    if index >= len(items):
        return ""
    item = items[index]
    if isinstance(item, dict):
        return str(item.get("text") or item.get("label") or "").strip()
    return str(item).strip()


def _contact_value(items: List[Any], kind: str) -> str:
    item = _contact_item(items, kind)
    if not isinstance(item, dict):
        return ""
    text = str(item.get("text") or item.get("label") or "").strip()
    if kind == "linkedin":
        return text or "LinkedIn"
    return text


def _contact_item(items: List[Any], kind: str) -> Dict[str, Any] | None:
    for item in items:
        if not isinstance(item, dict):
            continue
        item_kind = str(item.get("kind") or "").strip().lower()
        if item_kind == kind:
            return item

    for item in items:
        if not isinstance(item, dict):
            continue
        text = str(item.get("text") or item.get("label") or "").strip()
        label = str(item.get("label") or "").strip().lower()
        url = str(item.get("url") or "").strip().lower()
        if kind == "phone" and text and "mailto:" not in url and label not in {"linkedin", "github", "portfolio"} and text == str(item.get("label") or item.get("text") or "").strip():
            if any(ch.isdigit() for ch in text):
                return item
        if kind == "email" and (url.startswith("mailto:") or "@" in text):
            return item
        if kind == "address" and text and not any(marker in url for marker in ("mailto:", "http://", "https://")) and "@" not in text and not any(ch.isdigit() for ch in label):
            if label not in {"linkedin", "github", "portfolio"} and not any(ch.isdigit() for ch in label):
                return item
        if kind == "linkedin" and ("linkedin" in label or "linkedin.com" in url):
            return item
    return None


def _normalize_multiline_text(text: str) -> str:
    lines = [line.strip() for line in str(text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    kept: List[str] = []
    previous_blank = False
    for line in lines:
        is_blank = not line
        if is_blank and previous_blank:
            continue
        kept.append(line)
        previous_blank = is_blank
    return "\n".join(kept).strip()


def _normalize_bullet_text(text: str) -> str:
    normalized = _normalize_multiline_text(text)
    parts = [part.strip() for part in normalized.split("\n") if part.strip()]
    return " ".join(parts).strip()


def _markup_segments(text: str) -> List[tuple[str, bool]]:
    normalized = _normalize_multiline_text(text)
    pattern = re.compile(r"<b>(.*?)</b>", re.IGNORECASE | re.DOTALL)
    segments: List[tuple[str, bool]] = []
    last = 0
    for match in pattern.finditer(normalized):
        if match.start() > last:
            segments.append((normalized[last:match.start()], False))
        segments.append((match.group(1), True))
        last = match.end()
    if last < len(normalized):
        segments.append((normalized[last:], False))
    if not segments:
        segments.append((normalized, False))
    return [(part, bold) for part, bold in segments if part]


def _replace_tokens_in_document(doc: Document, replacements: Dict[str, str]) -> None:
    for paragraph in _iter_document_paragraphs(doc):
        _replace_tokens_in_paragraph(paragraph, replacements)


def _replace_tokens_in_paragraph(paragraph: Paragraph, replacements: Dict[str, str]) -> None:
    if not paragraph.runs:
        return

    for run in paragraph.runs:
        original = run.text or ""
        updated = original
        for src, dst in replacements.items():
            updated = updated.replace(src, dst)
        if updated != original:
            run.text = updated

    full_text = "".join(run.text or "" for run in paragraph.runs)
    if any(token in full_text for token in replacements):
        updated = full_text
        for src, dst in replacements.items():
            updated = updated.replace(src, dst)
        _set_paragraph_text(paragraph, updated)


def _set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    runs = list(paragraph.runs)
    if not runs:
        paragraph.add_run(text)
        return

    runs[0].text = text
    for run in runs[1:]:
        try:
            run._r.getparent().remove(run._r)
        except Exception:
            pass


def _clear_paragraph_content(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def _append_plain_run(paragraph: Paragraph, text: str) -> None:
    if not text:
        return
    paragraph.add_run(text)


def _add_hyperlink(paragraph: Paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _copy_run_style(source_run: Run, target_run: Run) -> None:
    source_rpr = source_run._r.rPr
    target_rpr = target_run._r.rPr
    if target_rpr is not None:
        target_run._r.remove(target_rpr)
    if source_rpr is not None:
        target_run._r.append(deepcopy(source_rpr))


def _style_source_run(paragraph: Paragraph) -> Run | None:
    runs = [run for run in paragraph.runs if (run.text or "").strip()]
    if not runs:
        runs = list(paragraph.runs)
    if not runs:
        return None
    return max(runs, key=lambda run: len((run.text or "").strip()))


def _set_paragraph_markup(paragraph: Paragraph, text: str) -> None:
    runs = list(paragraph.runs)
    if not runs:
        paragraph.add_run("")
        runs = list(paragraph.runs)

    first_run = runs[0]
    first_run.text = ""
    for run in runs[1:]:
        try:
            run._r.getparent().remove(run._r)
        except Exception:
            pass

    segments = _markup_segments(text)
    if not segments:
        return

    first_text, first_bold = segments[0]
    first_run.text = first_text
    first_run.bold = first_bold
    for segment_text, is_bold in segments[1:]:
        new_run = paragraph.add_run(segment_text)
        _copy_run_style(first_run, new_run)
        new_run.bold = is_bold


def _replace_markup_token_in_paragraph(paragraph: Paragraph, token: str, value: str) -> None:
    text = paragraph.text or ""
    rendered = text.replace(token, value or "")
    _set_paragraph_markup(paragraph, rendered)


def _render_literal_token_paragraph(paragraph: Paragraph, token: str, value: str) -> None:
    donor = _style_source_run(paragraph)
    runs = list(paragraph.runs)
    if not runs:
        paragraph.add_run(value)
        runs = list(paragraph.runs)
    if donor is not None and runs:
        _copy_run_style(donor, runs[0])
    _set_paragraph_text(paragraph, (paragraph.text or "").replace(token, value))


def _tighten_paragraph_spacing(paragraph: Paragraph) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = 0
    fmt.space_after = 0


def _insert_blank_paragraph_after(paragraph: Paragraph) -> Paragraph:
    blank_xml = OxmlElement("w:p")
    paragraph._p.addnext(blank_xml)
    blank = Paragraph(blank_xml, paragraph._parent)
    if blank.runs:
        _set_paragraph_text(blank, "")
    _tighten_paragraph_spacing(blank)
    return blank


def _iter_document_paragraphs(doc: Document) -> Iterable[Paragraph]:
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        yield from _iter_table_paragraphs(table)


def _iter_table_paragraphs(table: Table) -> Iterable[Paragraph]:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)


def _find_paragraph_containing(doc: Document, token: str) -> Paragraph | None:
    for paragraph in _iter_document_paragraphs(doc):
        if token in (paragraph.text or ""):
            return paragraph
    return None


def _find_table_with_any_text(doc: Document, tokens: Iterable[str]) -> Table | None:
    token_list = list(tokens)
    for table in doc.tables:
        for paragraph in _iter_table_paragraphs(table):
            text = paragraph.text or ""
            if any(token in text for token in token_list):
                return table
    return None


def _insert_paragraph_after(paragraph: Paragraph, source: Paragraph | None = None) -> Paragraph:
    template = source or paragraph
    new_p = deepcopy(template._p)
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def _insert_table_after(node_paragraph: Paragraph, table: Table) -> Table:
    new_tbl = deepcopy(table._tbl)
    node_paragraph._p.addnext(new_tbl)
    return Table(new_tbl, table._parent)


def _insert_paragraph_after_table(table: Table, source: Paragraph) -> Paragraph:
    new_p = deepcopy(source._p)
    table._tbl.addnext(new_p)
    return Paragraph(new_p, source._parent)


def _table_from_xml(table: Table, xml: str) -> Table:
    new_tbl = parse_xml(xml)
    return Table(new_tbl, table._parent)


def _paragraph_from_xml(paragraph: Paragraph, xml: str) -> Paragraph:
    new_p = parse_xml(xml)
    return Paragraph(new_p, paragraph._parent)


def _remove_paragraph(paragraph: Paragraph) -> None:
    try:
        element = paragraph._p
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)
    except Exception:
        pass


def _remove_table(table: Table) -> None:
    try:
        element = table._tbl
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)
    except Exception:
        pass


def _replace_tokens_in_table(table: Table, replacements: Dict[str, str]) -> None:
    for paragraph in _iter_table_paragraphs(table):
        _replace_tokens_in_paragraph(paragraph, replacements)


def _render_skill_block(doc: Document, skills: List[Dict[str, str]]) -> None:
    template = _find_paragraph_containing(doc, "[Category1]")
    if template is None:
        return

    template_xml = template._p.xml
    items = skills or [{"category": "", "details": ""}]
    cursor = template
    for idx, skill in enumerate(items):
        target = template if idx == 0 else _paragraph_from_xml(template, template_xml)
        if idx > 0:
            cursor._p.addnext(target._p)
        _replace_tokens_in_paragraph(
            target,
            {
                "[Category1]": skill.get("category") or "",
                "[Detail1]": skill.get("details") or "",
            },
        )
        cursor = target


def _render_experience_block(doc: Document, experiences: List[Dict[str, Any]]) -> None:
    template_table = _find_table_with_any_text(
        doc,
        ["[Role1]", "[Company1]", "[Company_adress1]", "[Date_range1]"],
    )
    items = experiences or [
        {
            "role": "",
            "company": "",
            "location": "",
            "duration": "",
            "bullets": [""],
        }
    ]

    if template_table is None:
        _render_experience_paragraph_block(doc, items)
        return

    template_bullet = _find_paragraph_containing(doc, "[Description1]")
    if template_bullet is None:
        return

    table_template_xml = template_table._tbl.xml
    bullet_template_xml = template_bullet._p.xml

    current_bullet_anchor = template_bullet
    for exp_idx, exp in enumerate(items):
        if exp_idx == 0:
            table = template_table
        else:
            table = _table_from_xml(template_table, table_template_xml)
            current_bullet_anchor._p.addnext(table._tbl)
        _replace_tokens_in_table(
            table,
            {
                "[Role1]": exp.get("role") or "",
                "[Company1]": exp.get("company") or "",
                "[Company_adress1]": exp.get("location") or "",
                "[Date_range1]": exp.get("duration") or "",
            },
        )

        bullets = exp.get("bullets") or [""]
        if exp_idx == 0:
            first_bullet = template_bullet
        else:
            first_bullet = _paragraph_from_xml(template_bullet, bullet_template_xml)
            table._tbl.addnext(first_bullet._p)
        _replace_markup_token_in_paragraph(first_bullet, "[Description1]", bullets[0] if bullets else "")
        _tighten_paragraph_spacing(first_bullet)

        bullet_cursor = first_bullet
        for bullet_text in bullets[1:]:
            next_bullet = _paragraph_from_xml(template_bullet, bullet_template_xml)
            bullet_cursor._p.addnext(next_bullet._p)
            _replace_markup_token_in_paragraph(next_bullet, "[Description1]", bullet_text)
            _tighten_paragraph_spacing(next_bullet)
            bullet_cursor = next_bullet

        current_bullet_anchor = bullet_cursor
        if exp_idx < len(items) - 1:
            current_bullet_anchor = _insert_blank_paragraph_after(bullet_cursor)


def _render_experience_paragraph_block(doc: Document, experiences: List[Dict[str, Any]]) -> None:
    role_template = _find_paragraph_containing(doc, "[Role1]") or _find_paragraph_containing(doc, "[Date_range1]")
    company_template = _find_paragraph_containing(doc, "[Company1]")
    duration_template = _find_paragraph_containing(doc, "[Date_range1]")
    bullet_template = _find_paragraph_containing(doc, "[Description1]")
    if role_template is None or bullet_template is None:
        return

    header_templates: List[Paragraph] = []
    for paragraph in [role_template, company_template, duration_template]:
        if paragraph is None:
            continue
        if any(existing._p is paragraph._p for existing in header_templates):
            continue
        header_templates.append(paragraph)
    if not header_templates:
        return

    header_template_xml = [paragraph._p.xml for paragraph in header_templates]
    bullet_template_xml = bullet_template._p.xml

    current_anchor = bullet_template
    for exp_idx, exp in enumerate(experiences):
        if exp_idx == 0:
            header_paragraphs = header_templates
            first_bullet = bullet_template
        else:
            header_paragraphs = []
            anchor = current_anchor
            for paragraph, paragraph_xml in zip(header_templates, header_template_xml):
                clone = _paragraph_from_xml(paragraph, paragraph_xml)
                anchor._p.addnext(clone._p)
                header_paragraphs.append(clone)
                anchor = clone

            first_bullet = _paragraph_from_xml(bullet_template, bullet_template_xml)
            header_paragraphs[-1]._p.addnext(first_bullet._p)

        replacements = {
            "[Role1]": exp.get("role") or "",
            "[Company1]": exp.get("company") or "",
            "[Company_adress1]": exp.get("location") or "",
            "[Date_range1]": exp.get("duration") or "",
        }
        for paragraph in header_paragraphs:
            _replace_tokens_in_paragraph(paragraph, replacements)

        bullets = exp.get("bullets") or [""]
        _replace_markup_token_in_paragraph(first_bullet, "[Description1]", bullets[0] if bullets else "")
        _tighten_paragraph_spacing(first_bullet)

        bullet_cursor = first_bullet
        for bullet_text in bullets[1:]:
            next_bullet = _paragraph_from_xml(bullet_template, bullet_template_xml)
            bullet_cursor._p.addnext(next_bullet._p)
            _replace_markup_token_in_paragraph(next_bullet, "[Description1]", bullet_text)
            _tighten_paragraph_spacing(next_bullet)
            bullet_cursor = next_bullet

        current_anchor = bullet_cursor
        if exp_idx < len(experiences) - 1:
            current_anchor = _insert_blank_paragraph_after(bullet_cursor)


def _render_education_block(doc: Document, education: List[Dict[str, str]]) -> None:
    school_template = _find_paragraph_containing(doc, "[University_name]")
    degree_template = _find_paragraph_containing(doc, "[Degree]")
    duration_template = _find_paragraph_containing(doc, "[Education_date_range]")
    if school_template is None or degree_template is None or duration_template is None:
        return

    items = education or [{"school": "", "degree": "", "duration": ""}]
    current_anchor = duration_template
    for idx, item in enumerate(items):
        school = school_template if idx == 0 else _insert_paragraph_after(current_anchor, school_template)
        _replace_tokens_in_paragraph(school, {"[University_name]": item.get("school") or ""})

        degree = degree_template if idx == 0 else _insert_paragraph_after(school, degree_template)
        _replace_tokens_in_paragraph(degree, {"[Degree]": item.get("degree") or ""})

        duration = duration_template if idx == 0 else _insert_paragraph_after(degree, duration_template)
        _replace_tokens_in_paragraph(
            duration,
            {"[Education_date_range]": item.get("duration") or ""},
        )
        current_anchor = duration


def _render_summary_block(doc: Document, summary: str) -> None:
    summary_paragraph = _find_paragraph_containing(doc, "[Summary]") or _find_paragraph_containing(doc, "[SUMMARY]")
    if summary_paragraph is None:
        return
    _set_paragraph_markup(summary_paragraph, summary)


def _render_name_block(doc: Document, name: str) -> None:
    name_paragraph = _find_paragraph_containing(doc, "[Name]") or _find_paragraph_containing(doc, "[NAME]")
    if name_paragraph is None:
        return
    _set_paragraph_markup(name_paragraph, f"<b>{(name or '').strip()}</b>")


def _render_heading_block(doc: Document, token: str, value: str) -> None:
    paragraph = _find_paragraph_containing(doc, token)
    if paragraph is None:
        return
    _render_literal_token_paragraph(paragraph, token, value)


def _render_contact_link_block(doc: Document, tokens: List[str], item: Dict[str, Any] | None) -> None:
    if item is None:
        return

    text = str(item.get("text") or item.get("label") or "").strip() or "LinkedI